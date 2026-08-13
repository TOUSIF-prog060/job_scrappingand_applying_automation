import os
import re
import json
import sys

def parse_pdf(file_path):
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    except Exception as e:
        sys.stderr.write(f"[PythonResumeParser] pdfplumber error: {e}, falling back to pypdf\n")

    if not text.strip():
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        except Exception as e:
            sys.stderr.write(f"[PythonResumeParser] pypdf error: {e}\n")

    return text

def parse_docx(file_path):
    text = ""
    try:
        import docx
        doc = docx.Document(file_path)
        for p in doc.paragraphs:
            if p.text:
                text += p.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text += " | ".join(row_text) + "\n"
    except Exception as e:
        sys.stderr.write(f"[PythonResumeParser] docx parse error: {e}\n")
    return text

def parse_txt(file_path):
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        sys.stderr.write(f"[PythonResumeParser] txt parse error: {e}\n")
        return ""

SKILL_DICTIONARY = [
    "JavaScript", "TypeScript", "React", "React.js", "Node.js", "Node", "Python",
    "Java", "C++", "C#", "Go", "Golang", "Rust", "Ruby", "PHP", "Swift", "Kotlin",
    "HTML", "CSS", "Tailwind", "TailwindCSS", "Sass", "Bootstrap",
    "SQL", "MySQL", "PostgreSQL", "MongoDB", "Redis", "SQLite", "DynamoDB", "Snowflake", "BigQuery",
    "Express", "Next.js", "Vite", "Django", "Flask", "Spring", "FastAPI", "Nuxt.js",
    "AWS", "GCP", "Azure", "Docker", "Kubernetes", "Terraform", "Ansible", "CI/CD",
    "Git", "GitHub", "Linux", "REST", "RESTful", "GraphQL", "gRPC", "WebSockets",
    "Playwright", "Puppeteer", "Selenium", "Jest", "Cypress", "PyTest",
    "LLM", "AI", "Machine Learning", "Deep Learning", "PyTorch", "TensorFlow", "OpenAI",
    "LangChain", "LlamaIndex", "RAG", "Vector DB", "Pinecone", "Milvus", "Chroma",
    "Pandas", "NumPy", "Scikit-Learn", "SciPy", "Keras", "Hugging Face", "Transformers",
    "Kafka", "Spark", "Airflow", "Microservices", "System Design", "OOP"
]

def extract_candidate_data(file_path, current_candidate=None):
    if current_candidate is None:
        current_candidate = {}

    if not os.path.exists(file_path):
        return current_candidate

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        text = parse_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        text = parse_docx(file_path)
    else:
        text = parse_txt(file_path)

    if not text.strip():
        return current_candidate

    extracted = dict(current_candidate)
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    # 1. Email
    email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    if email_match:
        extracted["email"] = email_match.group(0)

    # 2. Phone
    phone_match = re.search(r'(?:\+\d{1,3}[\s-]?)?\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4}|\+?\d{10,13}', text)
    if phone_match:
        extracted["phone"] = phone_match.group(0)

    # 3. LinkedIn & GitHub
    li_match = re.search(r'(?:https?:\/\/)?(?:www\.)?linkedin\.com\/in\/[a-zA-Z0-9_-]+\/?', text, re.IGNORECASE)
    if li_match:
        url = li_match.group(0)
        extracted["linkedin"] = url if url.startswith("http") else f"https://{url}"

    gh_match = re.search(r'(?:https?:\/\/)?(?:www\.)?github\.com\/[a-zA-Z0-9_-]+\/?', text, re.IGNORECASE)
    if gh_match:
        url = gh_match.group(0)
        extracted["github"] = url if url.startswith("http") else f"https://{url}"

    # 4. Name extraction
    possible_name = next((l for l in lines[:10] if len(l) < 50 and "@" not in l and "http" not in l and not re.search(r'\d{5,}', l) and "resume" not in l.lower()), None)
    if possible_name:
        clean_name = possible_name.split("|")[0].strip()
        parts = [p for p in re.sub(r'[^a-zA-Z\s]', '', clean_name).strip().split() if p]
        if len(parts) >= 2:
            extracted["firstName"] = parts[0].capitalize()
            extracted["lastName"] = " ".join([p.capitalize() for p in parts[1:]])
        elif len(parts) == 1:
            extracted["firstName"] = parts[0].capitalize()

    # 5. Location extraction
    loc_match = re.search(r'([A-Z][a-z]+(?:\s[A-Z][a-z]+)?),\s*([A-Z][a-z]+|[A-Z]{2})', text)
    if loc_match:
        extracted["location"] = loc_match.group(0)
    else:
        city_line = next((l for l in lines[:15] if re.search(r'bengaluru|bangalore|mumbai|delhi|pune|hyderabad|chennai|san francisco|new york|london|paris|remote', l, re.IGNORECASE)), None)
        if city_line:
            extracted["location"] = city_line.split("|")[0].strip()

    # 6. Skills extraction
    found_skills = set(extracted.get("skills", []))
    for skill in SKILL_DICTIONARY:
        escaped = re.escape(skill)
        if re.search(rf'\b{escaped}\b', text, re.IGNORECASE):
            found_skills.add(skill)
    if found_skills:
        extracted["skills"] = list(found_skills)

    # 7. Current Title / Role
    title_line = next((l for l in lines if re.search(r'engineer|developer|intern|manager|analyst|consultant|architect|specialist', l, re.IGNORECASE)), None)
    if title_line:
        extracted["currentTitle"] = title_line.split("|")[0].strip()[:60]

    # 8. Education Snippet
    edu_line = next((l for l in lines if re.search(r'b\.e\.|b\.tech|b\.s\.|m\.s\.|master|bachelor|degree|university|college', l, re.IGNORECASE)), None)
    if edu_line:
        extracted["education"] = edu_line[:100]

    return extracted

if __name__ == "__main__":
    if len(sys.argv) > 1:
        target_path = sys.argv[1]
        base_candidate = {}
        if len(sys.argv) > 2:
            try:
                base_candidate = json.loads(sys.argv[2])
            except Exception:
                pass
        result = extract_candidate_data(target_path, base_candidate)
        print(json.dumps(result, indent=2))
